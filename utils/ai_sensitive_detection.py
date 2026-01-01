"""
AI脱敏识别工具
使用千问API识别附件中的敏感信息
使用Qwen-VL识别图片，使用Qwen-Plus识别文本
"""
import os
import base64
import requests
from typing import List, Dict, Optional
from config import Config


class SensitiveDetector:
    """敏感信息检测器"""
    
    def __init__(self, api_key: str = None, api_base_url: str = None, vl_api_base_url: str = None):
        """
        初始化检测器
        
        Args:
            api_key: 千问API密钥
            api_base_url: 文本生成API基础URL（用于Qwen-Plus）
            vl_api_base_url: 视觉模型API基础URL（用于Qwen-VL）
        """
        self.api_key = api_key or os.environ.get('QWEN_API_KEY') or Config.QWEN_API_KEY
        self.api_base_url = api_base_url or os.environ.get('QWEN_API_BASE_URL') or getattr(Config, 'QWEN_API_BASE_URL', 'https://dashscope.aliyuncs.com/api/v1/services/aigc/text-generation/generation')
        self.vl_api_base_url = vl_api_base_url or os.environ.get('QWEN_VL_API_BASE_URL') or getattr(Config, 'QWEN_VL_API_BASE_URL', 'https://dashscope.aliyuncs.com/api/v1/services/aigc/multimodal-generation/generation')
        self.sensitive_keywords = getattr(Config, 'SENSITIVE_KEYWORDS', ['西南交通大学'])
    
    def encode_image(self, image_path: str) -> str:
        """将图片编码为base64"""
        with open(image_path, 'rb') as image_file:
            return base64.b64encode(image_file.read()).decode('utf-8')
    
    def detect_image(self, image_path: str) -> Dict:
        """
        检测图片中的敏感信息
        
        Args:
            image_path: 图片文件路径
            
        Returns:
            {
                'has_sensitive': bool,  # 是否包含敏感信息
                'detected_keywords': List[str],  # 检测到的关键词
                'details': str,  # 详细说明
                'error': str  # 错误信息（如果有）
            }
        """
        try:
            # 编码图片
            base64_image = self.encode_image(image_path)
            
            # 构建关键词提示
            keywords_str = '、'.join(self.sensitive_keywords)
            prompt = f"""你是一个专业的文档审核助手。请检查这张图片中是否包含以下敏感关键词：{keywords_str}。

要求：
1. 仔细识别图片中的所有文字内容（包括图片中的文字、水印、标题等）
2. 如果发现包含上述任一关键词，请用简洁、专业的语言详细说明
3. 说明要清晰易懂，便于管理员快速了解情况
4. 输出格式必须左对齐，每行从行首开始，不要使用缩进或空格

回复格式（如果包含敏感关键词，注意：每行从行首开始，左对齐）：
检测结果：包含敏感信息
发现的关键词：[列出所有发现的关键词，用顿号分隔]
详细位置：[用简洁的语言说明关键词在图片中的具体位置，如"标题区域"、"正文第一段"、"水印位置"等]

回复格式（如果不包含，注意：每行从行首开始，左对齐）：
检测结果：未发现敏感信息

注意：请只返回检测结果，不要添加任何解释性文字或额外说明。输出格式必须左对齐，每行从行首开始。"""
            
            # 调用千问视觉API (Qwen-VL)
            headers = {
                'Authorization': f'Bearer {self.api_key}',
                'Content-Type': 'application/json'
            }
            
            # 根据文件类型确定MIME类型
            image_ext = os.path.splitext(image_path)[1].lower()
            mime_type = 'image/jpeg'
            if image_ext == '.png':
                mime_type = 'image/png'
            elif image_ext == '.gif':
                mime_type = 'image/gif'
            elif image_ext == '.bmp':
                mime_type = 'image/bmp'
            elif image_ext == '.webp':
                mime_type = 'image/webp'
            
            payload = {
                "model": "qwen-vl-max",  # 使用Qwen-VL视觉模型
                "input": {
                    "messages": [
                        {
                            "role": "user",
                            "content": [
                                {
                                    "image": f"data:{mime_type};base64,{base64_image}"
                                },
                                {
                                    "text": prompt
                                }
                            ]
                        }
                    ]
                },
                "parameters": {
                    "temperature": 0.1,  # 降低温度以获得更确定的结果
                }
            }
            
            response = requests.post(self.vl_api_base_url, headers=headers, json=payload, timeout=60)
            response.raise_for_status()
            
            result = response.json()
            
            # 解析结果 - 支持多种可能的API返回格式
            content = None
            
            # 格式1: output.choices[0].message.content (标准格式)
            if 'output' in result and 'choices' in result['output'] and len(result['output']['choices']) > 0:
                message_content = result['output']['choices'][0].get('message', {}).get('content')
                
                # content可能是字符串、列表或字典
                if isinstance(message_content, str):
                    content = message_content
                elif isinstance(message_content, list):
                    # 如果是列表，尝试提取文本
                    text_parts = []
                    for item in message_content:
                        if isinstance(item, dict):
                            if 'text' in item:
                                text_parts.append(item['text'])
                            elif 'content' in item:
                                text_parts.append(item['content'])
                        elif isinstance(item, str):
                            text_parts.append(item)
                    content = '\n'.join(text_parts) if text_parts else str(message_content)
                elif isinstance(message_content, dict):
                    # 如果是字典，尝试提取text字段
                    content = message_content.get('text', str(message_content))
                else:
                    content = str(message_content) if message_content else None
            
            # 格式2: output.text (备选格式)
            if not content and 'output' in result and 'text' in result['output']:
                content = result['output']['text']
            
            # 格式3: choices[0].message.content (直接格式)
            if not content and 'choices' in result and len(result['choices']) > 0:
                message_content = result['choices'][0].get('message', {}).get('content')
                if isinstance(message_content, str):
                    content = message_content
                elif isinstance(message_content, list):
                    text_parts = []
                    for item in message_content:
                        if isinstance(item, dict) and 'text' in item:
                            text_parts.append(item['text'])
                        elif isinstance(item, str):
                            text_parts.append(item)
                    content = '\n'.join(text_parts) if text_parts else str(message_content)
            
            # 格式4: text (直接文本)
            if not content and 'text' in result:
                content = result['text']
            
            if content:
                # 确保content是字符串
                if not isinstance(content, str):
                    content = str(content)
                
                # 解析检测结果
                has_sensitive = '包含敏感信息' in content or '发现的关键词' in content
                detected_keywords = []
                
                if has_sensitive:
                    # 提取关键词
                    for keyword in self.sensitive_keywords:
                        if keyword in content:
                            detected_keywords.append(keyword)
                
                return {
                    'has_sensitive': has_sensitive,
                    'detected_keywords': detected_keywords,
                    'details': content,
                    'error': None
                }
            else:
                # 如果无法解析，返回原始结果用于调试
                import json
                error_details = f'API返回格式异常，实际返回：{json.dumps(result, ensure_ascii=False, indent=2)[:500]}'
                return {
                    'has_sensitive': False,
                    'detected_keywords': [],
                    'details': error_details,
                    'error': 'API返回格式异常'
                }
                
        except Exception as e:
            return {
                'has_sensitive': False,
                'detected_keywords': [],
                'details': f'识别失败：{str(e)}',
                'error': str(e)
            }
    
    def detect_pdf(self, file_path: str) -> Dict:
        """
        检测PDF文件中的敏感信息
        使用PyPDF2提取文本，然后使用Qwen-Plus进行识别
        
        Args:
            file_path: PDF文件路径
            
        Returns:
            检测结果字典
        """
        try:
            # 尝试使用PyPDF2提取PDF文本
            try:
                import PyPDF2
                
                text_content = ""
                with open(file_path, 'rb') as pdf_file:
                    pdf_reader = PyPDF2.PdfReader(pdf_file)
                    for page_num, page in enumerate(pdf_reader.pages, 1):
                        page_text = page.extract_text()
                        if page_text:
                            text_content += f"\n--- 第{page_num}页 ---\n{page_text}\n"
                
                if not text_content.strip():
                    return {
                        'has_sensitive': False,
                        'detected_keywords': [],
                        'details': 'PDF文件中没有提取到文本内容，可能是扫描版PDF或图片PDF。请手动检查。',
                        'error': '无法提取PDF文本内容'
                    }
                
                # 使用提取的文本内容进行识别
                return self.detect_text_content(text_content)
                
            except ImportError:
                return {
                    'has_sensitive': False,
                    'detected_keywords': [],
                    'details': '注意：PDF文件的识别需要安装PyPDF2库。请运行：pip install PyPDF2',
                    'error': 'PyPDF2库未安装'
                }
            except Exception as e:
                return {
                    'has_sensitive': False,
                    'detected_keywords': [],
                    'details': f'PDF文本提取失败：{str(e)}。请检查PDF文件是否损坏或需要密码。',
                    'error': f'PDF提取失败：{str(e)}'
                }
            
        except Exception as e:
            return {
                'has_sensitive': False,
                'detected_keywords': [],
                'details': f'识别失败：{str(e)}',
                'error': str(e)
            }
    
    def detect_text_content(self, text_content: str) -> Dict:
        """
        检测文本内容中的敏感信息（使用Qwen-Plus）
        
        Args:
            text_content: 文本内容
            
        Returns:
            检测结果字典
        """
        try:
            keywords_str = '、'.join(self.sensitive_keywords)
            prompt = f"""你是一个专业的文档审核助手。请检查以下文本内容中是否包含以下敏感关键词：{keywords_str}。

要求：
1. 仔细检查文本中的所有内容
2. 如果发现包含上述任一关键词，请用简洁、专业的语言详细说明
3. 说明要清晰易懂，便于管理员快速了解情况
4. 输出格式必须左对齐，每行从行首开始，不要使用缩进或空格

回复格式（如果包含敏感关键词，注意：每行从行首开始，左对齐）：
检测结果：包含敏感信息
发现的关键词：[列出所有发现的关键词，用顿号分隔]
详细位置：[用简洁的语言说明关键词在文本中的具体位置，如"第X段"、"第X行"、"标题"、"正文部分"等，如果出现多次请说明所有位置]

回复格式（如果不包含，注意：每行从行首开始，左对齐）：
检测结果：未发现敏感信息

注意：请只返回检测结果，不要添加任何解释性文字或额外说明。输出格式必须左对齐，每行从行首开始。

文本内容：
{text_content[:8000]}"""  # 限制文本长度
            
            headers = {
                'Authorization': f'Bearer {self.api_key}',
                'Content-Type': 'application/json'
            }
            
            payload = {
                "model": "qwen-plus",
                "input": {
                    "messages": [
                        {
                            "role": "user",
                            "content": prompt
                        }
                    ]
                },
                "parameters": {
                    "temperature": 0.1,
                }
            }
            
            response = requests.post(self.api_base_url, headers=headers, json=payload, timeout=60)
            response.raise_for_status()
            
            result = response.json()
            
            # 解析结果 - 支持多种可能的API返回格式
            content = None
            
            # 格式1: output.choices[0].message.content (标准格式)
            if 'output' in result and 'choices' in result['output'] and len(result['output']['choices']) > 0:
                if 'message' in result['output']['choices'][0] and 'content' in result['output']['choices'][0]['message']:
                    content = result['output']['choices'][0]['message']['content']
            
            # 格式2: output.text (备选格式)
            if not content and 'output' in result and 'text' in result['output']:
                content = result['output']['text']
            
            # 格式3: choices[0].message.content (直接格式)
            if not content and 'choices' in result and len(result['choices']) > 0:
                if 'message' in result['choices'][0] and 'content' in result['choices'][0]['message']:
                    content = result['choices'][0]['message']['content']
            
            # 格式4: text (直接文本)
            if not content and 'text' in result:
                content = result['text']
            
            if content:
                # 解析检测结果
                has_sensitive = '包含敏感信息' in content or '发现的关键词' in content
                detected_keywords = []
                
                if has_sensitive:
                    # 提取关键词
                    for keyword in self.sensitive_keywords:
                        if keyword in content:
                            detected_keywords.append(keyword)
                
                return {
                    'has_sensitive': has_sensitive,
                    'detected_keywords': detected_keywords,
                    'details': content,
                    'error': None
                }
            else:
                # 如果无法解析，返回原始结果用于调试
                import json
                error_details = f'API返回格式异常，实际返回：{json.dumps(result, ensure_ascii=False, indent=2)[:500]}'
                return {
                    'has_sensitive': False,
                    'detected_keywords': [],
                    'details': error_details,
                    'error': 'API返回格式异常'
                }
                
        except Exception as e:
            return {
                'has_sensitive': False,
                'detected_keywords': [],
                'details': f'识别失败：{str(e)}',
                'error': str(e)
            }
    
    def detect_word(self, file_path: str, file_type: str) -> Dict:
        """
        检测Word文件中的敏感信息
        使用python-docx提取DOCX文本，然后使用Qwen-Plus进行识别
        
        Args:
            file_path: Word文件路径
            file_type: 文件类型（doc或docx）
            
        Returns:
            检测结果字典
        """
        try:
            file_type_lower = file_type.lower()
            
            if file_type_lower == 'docx':
                # 使用python-docx解析docx文件
                try:
                    from docx import Document
                    
                    doc = Document(file_path)
                    text_content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                    
                    # 也提取表格内容
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = ' '.join([cell.text for cell in row.cells])
                            if row_text.strip():
                                text_content += '\n' + row_text
                    
                    if not text_content.strip():
                        return {
                            'has_sensitive': False,
                            'detected_keywords': [],
                            'details': 'Word文档中没有提取到文本内容，请手动检查。',
                            'error': '无法提取文档文本'
                        }
                    
                    # 使用提取的文本内容进行识别
                    return self.detect_text_content(text_content)
                    
                except ImportError:
                    return {
                        'has_sensitive': False,
                        'detected_keywords': [],
                        'details': '注意：DOCX文件的识别需要安装python-docx库。请运行：pip install python-docx',
                        'error': 'python-docx库未安装'
                    }
                except Exception as e:
                    return {
                        'has_sensitive': False,
                        'detected_keywords': [],
                        'details': f'DOCX文件解析失败：{str(e)}。请检查文件是否损坏。',
                        'error': f'DOCX解析失败：{str(e)}'
                    }
            
            elif file_type_lower == 'doc':
                # DOC格式需要其他库（如python-docx2txt、pypandoc或antiword）
                # 这里提供一个提示
                return {
                    'has_sensitive': False,
                    'detected_keywords': [],
                    'details': '注意：DOC格式（旧版Word）文件需要专门的解析库。建议安装python-docx2txt（pip install python-docx2txt）或pypandoc（pip install pypandoc）以支持DOC文件识别，或手动检查文件内容。',
                    'error': 'DOC格式需要额外解析库支持'
                }
            else:
                return {
                    'has_sensitive': False,
                    'detected_keywords': [],
                    'details': f'不支持的文件类型：{file_type}',
                    'error': f'不支持的文件类型：{file_type}'
                }
            
        except Exception as e:
            return {
                'has_sensitive': False,
                'detected_keywords': [],
                'details': f'识别失败：{str(e)}',
                'error': str(e)
            }
    
    def detect_attachment(self, file_path: str, file_type: str) -> Dict:
        """
        检测附件中的敏感信息（自动识别文件类型）
        
        Args:
            file_path: 文件路径
            file_type: 文件类型（pdf, jpg, png, doc, docx等）
            
        Returns:
            检测结果字典
        """
        file_type_lower = file_type.lower()
        
        # 图片类型
        if file_type_lower in ['jpg', 'jpeg', 'png', 'gif', 'bmp', 'webp']:
            return self.detect_image(file_path)
        # PDF文档
        elif file_type_lower == 'pdf':
            return self.detect_pdf(file_path)
        # Word文档
        elif file_type_lower in ['doc', 'docx']:
            return self.detect_word(file_path, file_type)
        else:
            return {
                'has_sensitive': False,
                'detected_keywords': [],
                'details': f'不支持的文件类型：{file_type}，当前仅支持图片（JPG、PNG、GIF等）、PDF和Word文档（DOC、DOCX）',
                'error': f'不支持的文件类型：{file_type}'
            }

